from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/Robinson-PrintWorks-90-Day-Growth-Plan.docx")

# compact_reference_guide preset tokens
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "56616F"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
CYAN = "0891B2"


def set_run_font(run, *, name="Calibri", size=11, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D7DEE8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering_definition(doc, *, bullet):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_list_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_bullet(doc, text, bullet_num_id, *, bold_lead=None):
    p = doc.add_paragraph()
    set_list_numbering(p, bullet_num_id)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_step(doc, text, decimal_num_id, *, bold_lead=None):
    p = doc.add_paragraph()
    set_list_numbering(p, decimal_num_id)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}: "), color=NAVY, bold=True)
    set_run_font(p.add_run(text), color=NAVY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_page_break(doc):
    doc.add_page_break()


def add_header_footer(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("ROBINSON PRINTWORKS"), size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("90-Day Growth & Accountability Plan"), size=8.5, color=MUTED)
    set_run_font(p.add_run("    |    "), size=8.5, color=MUTED)
    page_run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([fld_begin, instr, fld_sep, fld_end])
    set_run_font(page_run, size=8.5, color=MUTED)


def add_title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("BUSINESS GROWTH PLAN"), size=9.5, color=CYAN, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    set_run_font(title.add_run("Robinson PrintWorks"), size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(
        subtitle.add_run("A practical 90-day plan for growing the business and staying focused"),
        size=13,
        color=MUTED,
    )

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [4680, 4680])
    values = [
        ("PRIMARY FOCUS", "Functional local 3D-printing jobs"),
        ("TIME HORIZON", "Next 90 days"),
        ("CORE MARKET", "Rhode Island customers and partners"),
        ("OPERATING RHYTHM", "Weekly 3–2–1 system"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.rows[idx // 2].cells[idx % 2]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(label), size=8.5, color=BLUE, bold=True)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        set_run_font(p2.add_run(value), size=10.5, color=NAVY, bold=True)


def add_weekly_scorecard(doc):
    doc.add_heading("Weekly Scorecard", level=1)
    p = doc.add_paragraph(
        "Complete this once each Friday. The goal is to see whether activity is producing customers—not to create perfect bookkeeping."
    )
    p.paragraph_format.keep_with_next = True

    rows = [
        ("Qualified inquiries", "", ""),
        ("Quotes sent", "", ""),
        ("Jobs won", "", ""),
        ("Revenue", "", ""),
        ("Material and failed-print costs", "", ""),
        ("Reviews requested / received", "", ""),
        ("Outreach contacts", "", ""),
        ("Posts published", "", ""),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [4680, 2340, 2340])
    headers = ["Metric", "This Week", "This Month"]
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for metric, week, month in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((metric, week, month)):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value if value else " "), size=9.5)
    set_table_geometry(table, [4680, 2340, 2340])


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    add_header_footer(section)
    configure_styles(doc)
    bullet_num_id = add_numbering_definition(doc, bullet=True)
    minimum_week_bullet_num_id = add_numbering_definition(doc, bullet=True)
    decimal_num_id = add_numbering_definition(doc, bullet=False)
    checklist_num_id = add_numbering_definition(doc, bullet=False)

    add_title_block(doc)
    doc.add_heading("The 90-Day Goal", level=1)
    add_callout(doc, "FOCUS", "Build a repeatable pipeline for local functional 3D-printing jobs.")
    p = doc.add_paragraph(
        "The purpose of this plan is to keep Robinson PrintWorks moving forward without letting every new product idea, website tweak, or interesting model become a distraction. For the next 90 days, work should support at least one of four outcomes:"
    )
    p.paragraph_format.keep_with_next = True
    for text in (
        "Generate more qualified inquiries.",
        "Complete more profitable jobs.",
        "Build reviews and useful project examples.",
        "Make quoting and production easier to repeat.",
    ):
        add_bullet(doc, text, bullet_num_id)

    doc.add_heading("Primary Positioning", level=2)
    p = doc.add_paragraph()
    set_run_font(
        p.add_run("Practical Rhode Island 3D printing for replacement parts, prototypes, custom mounts, and small batches."),
        color=NAVY,
        bold=True,
    )
    p2 = doc.add_paragraph(
        "Gifts, figures, and fun prints can remain available, but functional work should lead the business because it answers clearer customer problems and creates better repeat and referral opportunities."
    )

    add_page_break(doc)
    doc.add_heading("The Weekly 3–2–1 System", level=1)
    for lead, text in (
        ("3 outreach actions. ", "Contact a repair shop, local business, school, maker group, property manager, or previous customer."),
        ("2 public posts. ", "Share a completed project, before-and-after, useful tip, or specific offer."),
        ("1 business improvement. ", "Publish a case study, improve quoting, update pricing, or create a reusable process."),
    ):
        add_bullet(doc, lead + text, bullet_num_id, bold_lead=lead)
    add_callout(doc, "RULE", "Complete the 3–2–1 commitments before spending time on optional redesigns, equipment research, or new product experiments.")

    doc.add_heading("Choose a Weekly Big 3", level=2)
    p = doc.add_paragraph("Every Sunday evening or Monday morning, select only three priorities. Example:")
    p.paragraph_format.keep_with_next = True
    for text in (
        "Publish the replacement vacuum-nozzle case study.",
        "Contact five appliance-repair businesses.",
        "Ask two completed customers for honest reviews.",
    ):
        add_step(doc, text, decimal_num_id)

    doc.add_heading("A Practical Weekly Schedule", level=1)
    schedule = [
        ("Monday · 30 min", "Choose the Big 3 and follow up on every open quote."),
        ("Tuesday · 60 min", "Work on one case study, service page, or product-photo set."),
        ("Wednesday · 30 min", "Contact five local prospects or potential partners."),
        ("Thursday · 30 min", "Prepare and publish one useful social post."),
        ("Friday · 20 min", "Update the scorecard, request reviews, and decide the next action for every lead."),
        ("Weekend", "Production, finished-project photography, and a second post when appropriate."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2600, 6760])
    for idx, value in enumerate(("Time Block", "Purpose")):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for block, purpose in schedule:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(block), size=9.5, color=DARK_BLUE, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(purpose), size=9.5)
    set_table_geometry(table, [2600, 6760])

    add_page_break(doc)
    doc.add_heading("The Minimum Viable Week", level=1)
    p = doc.add_paragraph(
        "Busy and exhausting weeks will happen. When the full plan is unrealistic, complete these four actions so the business still moves forward:"
    )
    p.paragraph_format.keep_with_next = True
    for text in (
        "Follow up with every open lead.",
        "Publish one completed-project photo.",
        "Contact one potential partner.",
        "Ask one satisfied customer for an honest review.",
    ):
        add_bullet(doc, text, minimum_week_bullet_num_id)

    doc.add_heading("Work-in-Progress Limits", level=1)
    p = doc.add_paragraph("At any one time, allow no more than:")
    p.paragraph_format.keep_with_next = True
    for text in (
        "One major website improvement.",
        "One marketing experiment.",
        "One new product idea.",
    ):
        add_bullet(doc, text, bullet_num_id)
    p = doc.add_paragraph(
        "Finish, measure, or intentionally stop one before starting another. Everything else goes into a Later list—not onto this week’s schedule."
    )

    doc.add_heading("The Lead List", level=1)
    p = doc.add_paragraph("Every potential customer should have these fields:")
    p.paragraph_format.keep_with_next = True
    for text in (
        "Name or business and contact information.",
        "What they need and the estimated value.",
        "Last contact date and the next action.",
        "A specific follow-up date.",
        "Status: new, quoting, waiting, won, lost, or completed.",
    ):
        add_bullet(doc, text, bullet_num_id)
    add_callout(doc, "NON-NEGOTIABLE", "No lead is allowed to sit in “waiting” without a dated follow-up action.")

    add_page_break(doc)
    doc.add_heading("Completed-Job Checklist", level=1)
    p = doc.add_paragraph(
        "Each finished job should create the raw material for future marketing and referrals. Use the same closing process every time:"
    )
    p.paragraph_format.keep_with_next = True
    steps = (
        "Confirm that the customer is satisfied.",
        "Photograph the finished part and, when possible, the installed result.",
        "Record the material, print time, price, costs, and lessons learned.",
        "Ask for an honest Google review—without offering an incentive.",
        "Ask permission to publish the project as a case study or post.",
        "Ask whether the customer knows anyone with a similar problem.",
        "Add the customer to the future follow-up list.",
    )
    for text in steps:
        add_step(doc, text, checklist_num_id)

    doc.add_heading("How We Can Review This Together", level=1)
    p = doc.add_paragraph(
        "A short weekly conversation can provide accountability without making either person feel responsible for managing the other. Keep it to 10–15 minutes."
    )
    for lead, text in (
        ("What moved? ", "Which actions produced a lead, quote, sale, review, or useful learning?"),
        ("What stalled? ", "Which commitment did not happen, and what made it difficult?"),
        ("What matters next? ", "What are the three priorities for the coming week?"),
        ("What should stop? ", "Is anything consuming time without helping the 90-day goal?"),
    ):
        add_bullet(doc, lead + text, bullet_num_id, bold_lead=lead)
    add_callout(doc, "ACCOUNTABILITY ROLE", "The goal is to ask clear questions and celebrate follow-through—not to become the business manager.")

    add_page_break(doc)
    add_weekly_scorecard(doc)
    doc.add_heading("Monthly Review", level=1)
    p = doc.add_paragraph("At the end of each month, answer:")
    p.paragraph_format.keep_with_next = True
    for text in (
        "What produced the best inquiries?",
        "Which jobs were most profitable?",
        "What repeatedly took too much time?",
        "What did customers repeatedly ask for?",
        "Which channel produced attention but no useful inquiries?",
        "What should be repeated, improved, or stopped next month?",
    ):
        add_bullet(doc, text, bullet_num_id)

    add_page_break(doc)
    doc.add_heading("Initial 90-Day Targets", level=1)
    targets = [
        ("Reviews", "10 genuine customer reviews"),
        ("Proof", "6 useful project case studies"),
        ("Demand", "10–20 qualified inquiries per month by the end of the period"),
        ("Service", "A consistent response time of one business day or less, when feasible"),
        ("Learning", "A clear view of the most profitable service and acquisition channel"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2200, 7160])
    for idx, value in enumerate(("Area", "Target")):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for area, target in targets:
        cells = table.add_row().cells
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(area), size=9.5, color=DARK_BLUE, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(target), size=9.5)
    set_table_geometry(table, [2200, 7160])

    doc.add_heading("First-Month Commitment", level=1)
    add_callout(
        doc,
        "START HERE",
        "Five outreach contacts per week, two public posts per week, three new case studies, and five review requests during the month.",
    )
    p = doc.add_paragraph(
        "Consistency is the objective. A modest plan completed every week will grow the business more reliably than an ambitious system used only occasionally."
    )
    p.paragraph_format.space_before = Pt(6)
    set_run_font(p.runs[0], color=NAVY, bold=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
